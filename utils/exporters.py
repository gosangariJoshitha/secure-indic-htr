"""
utils/exporters.py
===================
Turns a utils.layout_pipeline.LayoutDocument into each required export
format. Every exporter walks the SAME block list, so a paragraph, heading,
bullet/numbered list, blank line, or table renders consistently no matter
which format the user picks.
"""

from __future__ import annotations

import json as json_lib
import io
import datetime
import logging

from utils.layout_pipeline import LayoutDocument, BlockType, Alignment

logger = logging.getLogger("SecureDocAI.Exporters")


def _format_ocr_metadata(doc: LayoutDocument) -> str:
    engine = getattr(doc, "ocr_engine", "Custom OCR")
    lang = getattr(doc, "language", "Auto")
    score = getattr(doc, "final_ocr_score", None)
    if score is None:
        score = getattr(doc, "mean_confidence", 0.0) * 100.0
    t_proc = getattr(doc, "processing_time", 0.0)
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    return (
        f"--- SecureDocAI V2 OCR Metadata ---\n"
        f"OCR Engine : {engine}\n"
        f"Language   : {lang}\n"
        f"Date/Time  : {now}\n"
        f"Final OCR Score : {score:.2f}%\n"
        f"Processing Time : {t_proc:.2f}s\n"
        f"-----------------------------------"
    )


# ============================================================
# TXT — best effort: spacing preserved via indentation, tables as
# pipe-delimited rows.
# ============================================================
def export_txt(doc: LayoutDocument) -> str:
    lines = []
    for block in doc.blocks:
        if block.type == BlockType.BLANK:
            lines.append("")
            continue
        if block.type == BlockType.TABLE and block.table:
            lines.append(block.table.to_plain_text())
            lines.append("")
            continue

        prefix = ""
        if block.type == BlockType.BULLET_LIST:
            prefix = "• "
        for i, line in enumerate(block.lines):
            indent = "    " * line.indent_level
            marker = prefix
            if block.type == BlockType.NUMBERED_LIST:
                marker = f"{i + 1}. "
            lines.append(f"{indent}{marker}{line.text}")
            
    return "\n".join(lines)


# ============================================================
# Markdown — real headings, real lists, real tables.
# ============================================================
def export_markdown(doc: LayoutDocument) -> str:
    out = []
    for block in doc.blocks:
        if block.type == BlockType.BLANK:
            out.append("")
            continue
        if block.type == BlockType.TABLE and block.table:
            out.append(block.table.to_markdown())
            out.append("")
            continue
        if block.type == BlockType.HEADING:
            for line in block.lines:
                out.append(f"## {line.text}")
            continue
        if block.type == BlockType.BULLET_LIST:
            for line in block.lines:
                indent = "  " * line.indent_level
                out.append(f"{indent}- {line.text}")
            continue
        if block.type == BlockType.NUMBERED_LIST:
            for i, line in enumerate(block.lines):
                indent = "  " * line.indent_level
                out.append(f"{indent}{i + 1}. {line.text}")
            continue

        for line in block.lines:
            indent = "&nbsp;&nbsp;&nbsp;&nbsp;" * line.indent_level
            if line.alignment == Alignment.CENTER:
                out.append(f'<p align="center">{indent}{line.text}</p>')
            elif line.alignment == Alignment.RIGHT:
                out.append(f'<p align="right">{indent}{line.text}</p>')
            else:
                out.append(f"{indent}{line.text}")
        out.append("")
        
    return "\n".join(out)


# ============================================================
# HTML — semantic structure with inline alignment/indent, dark mode, responsive styles.
# ============================================================
def export_html(doc: LayoutDocument, title: str = "Digitized Document", watermark: str | None = None) -> str:
    body_parts = []
    for block in doc.blocks:
        if block.type == BlockType.BLANK:
            body_parts.append('<div class="sd-blank-line"></div>')
            continue
        if block.type == BlockType.TABLE and block.table:
            body_parts.append(_table_to_html(block.table))
            continue
        if block.type == BlockType.HEADING:
            for line in block.lines:
                body_parts.append(f'<h2 style="text-align:{line.alignment.value};">{_esc(line.text)}</h2>')
            continue
        if block.type == BlockType.BULLET_LIST:
            items = "".join(f"<li>{_esc(l.text)}</li>" for l in block.lines)
            body_parts.append(f"<ul>{items}</ul>")
            continue
        if block.type == BlockType.NUMBERED_LIST:
            items = "".join(f"<li>{_esc(l.text)}</li>" for l in block.lines)
            body_parts.append(f"<ol>{items}</ol>")
            continue

        for line in block.lines:
            indent_px = line.indent_level * 32
            body_parts.append(
                f'<p style="text-align:{line.alignment.value}; margin-left:{indent_px}px;">{_esc(line.text)}</p>'
            )

    body = "\n".join(body_parts)
    
    # Append Metadata container
    meta_html = _esc(_format_ocr_metadata(doc)).replace("\n", "<br>")
    body += f'\n<div class="sd-metadata">{meta_html}</div>'

    # Watermark CSS injection
    watermark_style = ""
    if watermark:
        watermark_style = f"""
  body::after {{
      content: "{_esc(watermark)}";
      position: fixed; top: 50%; left: 50%;
      transform: translate(-50%, -50%) rotate(-45deg);
      font-size: 80px; color: rgba(226, 232, 240, 0.4);
      pointer-events: none; z-index: -1;
  }}
"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{_esc(title)}</title>
<style>
  body {{ font-family: 'Noto Sans Telugu', 'Noto Sans Devanagari', Arial, sans-serif;
          max-width: 800px; margin: 40px auto; color: #0F172A; line-height: 1.7; padding: 0 16px; }}
  h2 {{ font-weight: 700; margin-top: 28px; }}
  p {{ margin: 6px 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
  td, th {{ border: 1px solid #94A3B8; padding: 8px 12px; text-align: left; }}
  .sd-blank-line {{ height: 18px; }}
  .sd-metadata {{ margin-top: 40px; padding: 12px; background: #F1F5F9; border-left: 4px solid #10B981; font-family: monospace; font-size: 13px; }}
  {watermark_style}
</style>
</head>
<body>
{body}
</body>
</html>"""


def _table_to_html(table) -> str:
    rows_html = []
    for row in table.cells:
        cells_html = "".join(f"<td>{_esc(c)}</td>" for c in row)
        rows_html.append(f"<tr>{cells_html}</tr>")
    return "<table>" + "".join(rows_html) + "</table>"


def _esc(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))


# ============================================================
# DOCX — python-docx: real headers, footers, watermarks.
# ============================================================
def export_docx(doc: LayoutDocument, title: str = "Digitized Document", watermark: str | None = None) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    align_map = {
        Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    }

    docx_doc = DocxDocument()
    docx_doc.add_heading(title, level=1)

    # Watermark header injection
    if watermark:
        header = docx_doc.sections[0].header
        hp = header.paragraphs[0]
        hp.text = f"[{watermark}] "
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_table_borders(table):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "94A3B8")
            borders.append(el)
        tbl_pr.append(borders)

    for block in doc.blocks:
        if block.type == BlockType.BLANK:
            docx_doc.add_paragraph("")
            continue

        if block.type == BlockType.TABLE and block.table:
            cells = block.table.cells
            if not cells:
                continue
            n_rows, n_cols = len(cells), max(len(r) for r in cells)
            table = docx_doc.add_table(rows=n_rows, cols=n_cols)
            table.style = "Table Grid"
            add_table_borders(table)
            for r, row in enumerate(cells):
                for c in range(n_cols):
                    table.cell(r, c).text = row[c] if c < len(row) else ""
            docx_doc.add_paragraph("")
            continue

        if block.type == BlockType.HEADING:
            for line in block.lines:
                docx_doc.add_heading(line.text, level=2)
            continue

        if block.type == BlockType.BULLET_LIST:
            for line in block.lines:
                p = docx_doc.add_paragraph(line.text, style="List Bullet")
                p.paragraph_format.left_indent = Pt(18 * (line.indent_level + 1))
            continue

        if block.type == BlockType.NUMBERED_LIST:
            for line in block.lines:
                p = docx_doc.add_paragraph(line.text, style="List Number")
                p.paragraph_format.left_indent = Pt(18 * (line.indent_level + 1))
            continue

        for line in block.lines:
            p = docx_doc.add_paragraph(line.text)
            p.alignment = align_map.get(line.alignment, WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.left_indent = Pt(18 * line.indent_level)

    # Append Metadata block
    docx_doc.add_paragraph("\n" + _format_ocr_metadata(doc))

    buf = io.BytesIO()
    docx_doc.save(buf)
    return buf.getvalue()


# ============================================================
# PDF — reportlab: watermarks & password protection.
# ============================================================
def export_pdf(doc: LayoutDocument, title: str = "Digitized Document", 
               watermark: str | None = None, password: str | None = None,
               doc_hash: str | None = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Table as RLTable, TableStyle,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

    buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buf, pagesize=A4,
                                 topMargin=20 * mm, bottomMargin=20 * mm,
                                 leftMargin=20 * mm, rightMargin=20 * mm)

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("SDBody", parent=styles["Normal"], fontSize=11, leading=16)
    heading_style = ParagraphStyle("SDHeading", parent=styles["Heading2"], fontSize=14, spaceAfter=8)
    title_style = ParagraphStyle("SDTitle", parent=styles["Heading1"], fontSize=18, spaceAfter=14)

    align_map = {Alignment.LEFT: TA_LEFT, Alignment.CENTER: TA_CENTER, Alignment.RIGHT: TA_RIGHT}

    flow = [RLParagraph(title, title_style), Spacer(1, 6)]

    for block in doc.blocks:
        if block.type == BlockType.BLANK:
            flow.append(Spacer(1, 12))
            continue

        if block.type == BlockType.TABLE and block.table:
            cells = block.table.cells
            if cells:
                rl_table = RLTable(cells)
                rl_table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#94A3B8")),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]))
                flow.append(rl_table)
                flow.append(Spacer(1, 10))
            continue

        if block.type == BlockType.HEADING:
            for line in block.lines:
                flow.append(RLParagraph(line.text, heading_style))
            continue

        if block.type in (BlockType.BULLET_LIST, BlockType.NUMBERED_LIST):
            for i, line in enumerate(block.lines):
                marker = "•" if block.type == BlockType.BULLET_LIST else f"{i + 1}."
                indent = 14 * (line.indent_level + 1)
                style = ParagraphStyle(f"list_{i}_{line.y}", parent=body_style, leftIndent=indent)
                flow.append(RLParagraph(f"{marker} {line.text}", style))
            continue

        for line in block.lines:
            style = ParagraphStyle(
                f"p_{line.y}", parent=body_style,
                alignment=align_map.get(line.alignment, TA_LEFT),
                leftIndent=14 * line.indent_level,
            )
            flow.append(RLParagraph(line.text or "&nbsp;", style))

    # Append Metadata Block
    flow.append(Spacer(1, 20))
    meta_style = ParagraphStyle("SDMeta", parent=body_style, fontSize=9, leading=13, textColor=colors.HexColor("#64748B"))
    flow.append(RLParagraph(_format_ocr_metadata(doc).replace("\n", "<br/>"), meta_style))

    # Append verification QR code
    h = doc_hash or getattr(doc, "integrity_hash", None)
    if h:
        try:
            import qrcode
            from reportlab.platypus import Image as RLImage
            qr = qrcode.QRCode(version=1, box_size=3, border=1)
            verify_url = f"http://localhost:8501/?verify={h}"
            qr.add_data(verify_url)
            qr.make(fit=True)
            img_qr = qr.make_image(fill_color="black", back_color="white")
            qr_buf = io.BytesIO()
            img_qr.save(qr_buf, format="PNG")
            qr_buf.seek(0)
            
            flow.append(Spacer(1, 15))
            flow.append(RLParagraph("<b>Ledger Verification QR Code:</b>", ParagraphStyle("QRTitle", parent=body_style, fontSize=9, textColor=colors.HexColor("#64748B"))))
            flow.append(Spacer(1, 4))
            flow.append(RLImage(qr_buf, width=60, height=60, hAlign="LEFT"))
        except Exception as e:
            logger.warning(f"Could not append QR verification code to PDF: {e}")

    # Watermark canvas helper
    def draw_watermark(canvas, doc_obj):
        wm_text = watermark
        if not wm_text:
            try:
                import streamlit as st
                import datetime
                email = st.session_state.get("user", {}).get("email")
                if email:
                    wm_text = f"SecureDocAI | {email} | {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
            except Exception:
                pass
                
        if wm_text:
            canvas.saveState()
            canvas.setFont('Helvetica-Bold', 28)
            canvas.setFillColor(colors.HexColor("#CBD5E1"))
            canvas.translate(297, 420)
            canvas.rotate(45)
            canvas.drawCentredString(0, 0, wm_text)
            canvas.restoreState()

    pdf_doc.build(flow, onFirstPage=draw_watermark, onLaterPages=draw_watermark)
    pdf_bytes = buf.getvalue()

    # Apply password protection if specified
    if password:
        pdf_bytes = _encrypt_pdf(pdf_bytes, password)

    return pdf_bytes


def _encrypt_pdf(pdf_bytes: bytes, password: str) -> bytes:
    """Uses pypdf or PyPDF2 to encrypt PDF bytes."""
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        try:
            from PyPDF2 import PdfReader, PdfWriter
        except ImportError:
            logger.warning("Neither pypdf nor PyPDF2 is installed; skipping PDF encryption.")
            return pdf_bytes

    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    
    out_buf = io.BytesIO()
    writer.write(out_buf)
    return out_buf.getvalue()


# ============================================================
# JSON — full layout metadata.
# ============================================================
def export_json(doc: LayoutDocument) -> str:
    payload = {
        "page_width": doc.page_width,
        "page_height": doc.page_height,
        "mean_confidence": round(doc.mean_confidence, 4),
        "char_count": doc.char_count,
        "ocr_version": "SecureDocAI V2",
        "ocr_engine": getattr(doc, "ocr_engine", "Custom OCR"),
        "language": getattr(doc, "language", "Auto"),
        "processing_time": getattr(doc, "processing_time", 0.0),
        "blocks": [],
    }

    for block in doc.blocks:
        block_data = {"type": block.type.value}
        if block.type == BlockType.TABLE and block.table:
            block_data["table"] = {
                "x": block.table.x, "y": block.table.y,
                "w": block.table.w, "h": block.table.h,
                "row_bounds": block.table.row_bounds,
                "col_bounds": block.table.col_bounds,
                "cells": block.table.cells,
            }
        else:
            block_data["lines"] = [
                {
                    "text": line.text,
                    "alignment": line.alignment.value,
                    "indent_level": line.indent_level,
                    "y": line.y,
                    "confidence": round(line.confidence, 4),
                }
                for line in block.lines
            ]
        payload["blocks"].append(block_data)

    return json_lib.dumps(payload, ensure_ascii=False, indent=2)


# ============================================================
# XML — structured schema for dataset annotations
# ============================================================
def export_xml(doc: LayoutDocument) -> str:
    """Exports LayoutDocument to XML representation for OCR evaluation and annotation."""
    import xml.etree.ElementTree as ET
    from xml.dom import minidom

    root = ET.Element("document")
    root.set("width", str(doc.page_width))
    root.set("height", str(doc.page_height))
    root.set("confidence", str(round(doc.mean_confidence, 4)))
    root.set("ocr_version", "SecureDocAI V2")
    root.set("language", getattr(doc, "language", "Auto"))
    
    for b_idx, block in enumerate(doc.blocks):
        b_el = ET.SubElement(root, "block")
        b_el.set("id", str(b_idx))
        b_el.set("type", block.type.value)
        
        if block.type == BlockType.TABLE and block.table:
            t_el = ET.SubElement(b_el, "table")
            t_el.set("x", str(block.table.x))
            t_el.set("y", str(block.table.y))
            t_el.set("w", str(block.table.w))
            t_el.set("h", str(block.table.h))
            for r_idx, row in enumerate(block.table.cells):
                r_el = ET.SubElement(t_el, "row")
                r_el.set("id", str(r_idx))
                for c_idx, cell in enumerate(row):
                    c_el = ET.SubElement(r_el, "cell")
                    c_el.set("id", str(c_idx))
                    c_el.text = cell
        else:
            for l_idx, line in enumerate(block.lines):
                l_el = ET.SubElement(b_el, "line")
                l_el.set("id", str(l_idx))
                l_el.set("y", str(line.y))
                l_el.set("alignment", line.alignment.value)
                l_el.set("indent", str(line.indent_level))
                l_el.set("confidence", str(round(line.confidence, 4)))
                l_el.text = line.text
                
    xmlstr = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")
    return xmlstr


# ============================================================
# Excel — pandas export of worksheets
# ============================================================
def export_excel(doc: LayoutDocument) -> bytes:
    """Exports all tables inside the LayoutDocument into an Excel workbook (.xlsx)."""
    try:
        import pandas as pd
    except ImportError:
        logger.warning("pandas is not installed; Excel export will return empty workbook.")
        return b""

    out_buf = io.BytesIO()
    try:
        with pd.ExcelWriter(out_buf, engine='openpyxl') as writer:
            table_idx = 1
            for block in doc.blocks:
                if block.type == BlockType.TABLE and block.table:
                    cells = block.table.cells
                    if cells:
                        df = pd.DataFrame(cells)
                        df.to_excel(writer, sheet_name=f"Table {table_idx}", index=False, header=False)
                        table_idx += 1
            if table_idx == 1:
                df = pd.DataFrame([["No tables found in this document."]])
                df.to_excel(writer, sheet_name="Summary", index=False, header=False)
    except Exception as e:
        logger.error(f"Excel export failed: {e}")
        return b""
            
    return out_buf.getvalue()


# ============================================================
# CSV — simple table outputs
# ============================================================
def export_csv(doc: LayoutDocument) -> str:
    """Exports all table blocks inside the LayoutDocument as CSV formatted string."""
    import csv
    out = io.StringIO()
    writer = csv.writer(out)
    table_idx = 1
    for block in doc.blocks:
        if block.type == BlockType.TABLE and block.table:
            cells = block.table.cells
            if cells:
                writer.writerow([f"--- Table {table_idx} ---"])
                writer.writerows(cells)
                writer.writerow([])
                table_idx += 1
    return out.getvalue()


# ============================================================
# ZIP — Single zip compilation containing all targets
# ============================================================
def export_zip(doc: LayoutDocument, title: str = "Digitized Document", 
               watermark: str | None = None, password: str | None = None) -> bytes:
    """Compiles all format exports, including V2 metadata descriptions, into a ZIP file."""
    import zipfile
    exports = export_all(doc, title, watermark=watermark, password=password)
    
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"{title}.txt", exports["txt"])
        zip_file.writestr(f"{title}.md", exports["md"])
        zip_file.writestr(f"{title}.html", exports["html"])
        zip_file.writestr(f"{title}.docx", exports["docx"])
        zip_file.writestr(f"{title}.pdf", exports["pdf"])
        zip_file.writestr(f"{title}.json", exports["json"])
        zip_file.writestr(f"{title}.xml", exports["xml"].encode("utf-8"))
        zip_file.writestr(f"{title}.csv", exports["csv"].encode("utf-8"))
        
        excel_bytes = exports.get("xlsx")
        if excel_bytes:
            zip_file.writestr(f"{title}.xlsx", excel_bytes)
            
        # Add metadata manifest descriptor
        manifest = {
            "title": title,
            "uuid": str(hash(title)),
            "created_time": datetime.datetime.now().isoformat(),
            "engine": getattr(doc, "ocr_engine", "Custom OCR"),
            "language": getattr(doc, "language", "Auto"),
            "confidence": getattr(doc, "mean_confidence", 0.0),
            "processing_time": getattr(doc, "processing_time", 0.0),
            "version": "SecureDocAI V2"
        }
        zip_file.writestr("Metadata.json", json_lib.dumps(manifest, indent=2))
        
    return zip_buf.getvalue()


# ============================================================
# Convenience: export to all formats at once, returns a dict of
# format_name -> bytes.
# ============================================================
def export_all(doc: LayoutDocument, title: str = "Digitized Document", 
               watermark: str | None = None, password: str | None = None) -> dict:
    result = {
        "txt": export_txt(doc).encode("utf-8"),
        "md": export_markdown(doc).encode("utf-8"),
        "html": export_html(doc, title, watermark=watermark).encode("utf-8"),
        "docx": export_docx(doc, title, watermark=watermark),
        "pdf": export_pdf(doc, title, watermark=watermark, password=password),
        "json": export_json(doc).encode("utf-8"),
        "xml": export_xml(doc),
        "csv": export_csv(doc),
    }
    
    excel_bytes = export_excel(doc)
    if excel_bytes:
         result["xlsx"] = excel_bytes
         
    return result
